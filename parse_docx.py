import docx
doc = docx.Document('agents/cam-agent/trinetra_cam_template.docx')
for p in doc.paragraphs:
    if "for" in p.text: print(repr(p.text))
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if "for" in r.text or "%" in r.text:
                        print(f"RUN: {repr(r.text)}")
